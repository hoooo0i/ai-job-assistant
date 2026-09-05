from __future__ import annotations

from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from src.privacy import redact_sensitive_info
from src.schemas import (
    CoverLetterDraft,
    JobComparisonItem,
    JobProfile,
    MatchAnalysis,
    RequirementMatch,
    ResumeProfile,
    ScoreResult,
)


STATUS_LABELS = {
    "matched": "已匹配",
    "partial": "部分匹配",
    "missing": "缺失",
    "unknown": "待确认",
}

APPLICATION_STATUS_LABELS = {
    "not_started": "尚未开始",
    "preparing": "准备材料",
    "applied": "已投递",
    "assessment": "笔试或测评",
    "interview": "面试中",
    "rejected": "未通过",
    "offer": "已获 Offer",
    "withdrawn": "已放弃",
}

PDF_FONT_NAME = "NotoSansSC"
PDF_FONT_PATH = Path(__file__).resolve().parents[1] / "assets" / "fonts" / "NotoSansSC.ttf"


def _score_text(score: ScoreResult) -> str:
    return "暂不可计算" if score.match_score is None else f"{score.match_score:.1f}%"


def _evidence_lines(match: RequirementMatch) -> list[str]:
    if match.evidence:
        return [
            f"{'[用户确认]' if item.source.value == 'user_confirmed' else '[简历]'} {item.text}"
            for item in match.evidence
        ]
    return [f"[简历] {text}" for text in match.resume_evidence]


def _suggestion_source(analysis: MatchAnalysis, original_text: str) -> str:
    lowered = original_text.casefold()
    for match in analysis.matches:
        for evidence in match.evidence:
            if evidence.source.value == "user_confirmed" and lowered in evidence.text.casefold():
                return "用户确认事实"
    return "简历原文"


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _style_docx_table_cell(cell, *, header: bool = False, alternate: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", 90), ("left", 100), ("bottom", 90), ("right", 100)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), "D9D9D9")
    if header:
        _set_cell_shading(cell, "23415C")
    elif alternate:
        _set_cell_shading(cell, "F2F6FA")
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            _set_docx_font(run)
            run.font.size = Pt(9.5)
            if header:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True


def _set_docx_font(run, name: str = "Noto Sans SC", east_asia: str = "Noto Sans SC") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def _add_docx_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(redact_sensitive_info(text))
    _set_docx_font(run)


def _style_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Noto Sans SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = document.styles["Title"]
    title.font.name = "Noto Sans SC"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans SC")
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.font.size = Pt(24)
    title_properties = title._element.get_or_add_pPr()
    title_border = title_properties.find(qn("w:pBdr"))
    if title_border is not None:
        title_properties.remove(title_border)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 12)]:
        style = document.styles[style_name]
        style.font.name = "Noto Sans SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans SC")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(size)
        style.font.bold = True


def build_docx_report(
    resume_profile: ResumeProfile,
    job_profile: JobProfile,
    analysis: MatchAnalysis,
    score: ScoreResult,
    interview_feedback: dict[str, dict] | None = None,
) -> bytes:
    """Create an editable, in-memory Word report without the raw resume."""
    _ = resume_profile
    document = Document()
    _style_docx(document)
    document.core_properties.title = "AI 求职分析报告"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_font(title.add_run("AI 求职分析报告"))
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_run = intro.add_run(f"{job_profile.company} · {job_profile.title}")
    intro_run.bold = True
    _set_docx_font(intro_run)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_font(
        meta.add_run(
            f"工作地点：{job_profile.location or '未提供'}    岗位类型：{job_profile.job_type or '未提供'}"
        )
    )

    document.add_heading("核心结果", level=1)
    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ["证据匹配度", "信息完整度", "评分版本"]
    values = [_score_text(score), f"{score.information_completeness:.1f}%", score.calculation_version]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        _set_cell_shading(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            _set_docx_font(run)
    for index, value in enumerate(values):
        cell = table.rows[1].cells[index]
        cell.text = value
        for run in cell.paragraphs[0].runs:
            _set_docx_font(run)

    requirements = {item.id: item for item in job_profile.requirements}
    document.add_heading("岗位要求与证据", level=1)
    for item in analysis.matches:
        requirement = requirements[item.requirement_id]
        document.add_heading(
            f"{requirement.normalized_name}  {STATUS_LABELS[item.status.value]}",
            level=2,
        )
        _add_docx_bullet(document, f"JD：{requirement.original_text}")
        _add_docx_bullet(document, f"判断：{item.explanation}")
        evidence_text = "；".join(_evidence_lines(item)) or "无可验证证据"
        _add_docx_bullet(document, f"证据：{evidence_text}")

    document.add_heading("简历优化建议", level=1)
    if not analysis.resume_suggestions:
        document.add_paragraph("本次没有通过证据校验的建议。")
    for suggestion in analysis.resume_suggestions:
        _add_docx_bullet(
            document,
            f"{_suggestion_source(analysis, suggestion.original_text)}：{suggestion.original_text}",
        )
        _add_docx_bullet(document, f"建议：{suggestion.suggested_text}")
        _add_docx_bullet(document, f"理由：{suggestion.reason}")

    document.add_heading("面试准备问题", level=1)
    if not analysis.interview_questions:
        document.add_paragraph("本次没有生成可用的面试问题。")
    for question in analysis.interview_questions:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(question.question)
        run.bold = True
        _set_docx_font(run)
        for point in question.answer_outline:
            _add_docx_bullet(document, point)

    feedback_items = interview_feedback or {}
    if feedback_items:
        document.add_heading("模拟面试复盘", level=1)
        for record in feedback_items.values():
            document.add_heading(record.get("question", "面试问题"), level=2)
            feedback = record.get("feedback", {})
            document.add_paragraph(
                "完整性 {}/5    STAR {}/5    岗位相关性 {}/5    表达清晰度 {}/5".format(
                    feedback.get("completeness_score", 0),
                    feedback.get("star_score", 0),
                    feedback.get("relevance_score", 0),
                    feedback.get("clarity_score", 0),
                )
            )
            for improvement in feedback.get("improvements", []):
                _add_docx_bullet(document, f"改进：{improvement}")

    document.add_heading("使用说明", level=1)
    document.add_paragraph(
        "本报告基于用户提供的信息生成，不代表录取概率或 ATS 通过率。"
        "报告不包含原始简历全文，常见联系方式已脱敏。"
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _comparison_match_text(item: JobComparisonItem) -> str:
    return "--" if item.match_score is None else f"{item.match_score:.1f}%"


def _comparison_stage_text(item: JobComparisonItem) -> str:
    analysis_stage = "最终分析" if item.stage == "final" else "待补充"
    application_stage = APPLICATION_STATUS_LABELS.get(
        item.application_status, "尚未开始"
    )
    return f"{analysis_stage}\n{application_stage}"


def build_job_comparison_docx(items: list[JobComparisonItem]) -> bytes:
    """Create an editable multi-job comparison report without resume source text."""
    document = Document()
    _style_docx(document)
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    normal = document.styles["Normal"]
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    document.core_properties.title = "多岗位求职对比报告"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_font(title.add_run("多岗位求职对比报告"))
    intro = document.add_paragraph(
        "本报告汇总当前简历在各岗位下的本地计算结果，便于安排查看和投递顺序。"
        "推荐值不是录取概率，仍需结合岗位要求和逐项证据判断。"
    )
    for run in intro.runs:
        _set_docx_font(run)

    document.add_heading("岗位排序", level=1)
    table = document.add_table(rows=1, cols=7)
    table.autofit = False
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    headers = ["序号", "公司与岗位", "分析与投递", "匹配度", "完整度", "ATS", "推荐值"]
    widths = [0.45, 2.45, 0.9, 0.7, 0.75, 0.55, 0.7]
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        cell.width = Inches(width)
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_docx_table_cell(cell, header=True)
    for rank, item in enumerate(items, start=1):
        row = table.add_row()
        values = [
            str(rank),
            f"{redact_sensitive_info(item.company)}\n{redact_sensitive_info(item.title)}",
            _comparison_stage_text(item),
            _comparison_match_text(item),
            f"{item.information_completeness:.1f}%",
            str(item.ats_score),
            f"{item.recommendation_score:.1f}",
        ]
        for index, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
            cell.width = Inches(width)
            cell.text = value
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 1 else WD_ALIGN_PARAGRAPH.CENTER
            )
            _style_docx_table_cell(cell, alternate=rank % 2 == 0)
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
        for cell in table.columns[index].cells:
            cell.width = Inches(width)

    document.add_heading("风险明细", level=1)
    if not items:
        document.add_paragraph("当前没有可比较的岗位。")
    for rank, item in enumerate(items, start=1):
        paragraph = document.add_paragraph(
            redact_sensitive_info(
                f"{rank}. {item.company} · {item.title}：硬性风险 {item.hard_risks} 项，"
                f"必须项缺口 {item.must_have_gaps} 项。"
            )
        )
        for run in paragraph.runs:
            _set_docx_font(run)

    document.add_heading("隐私与使用说明", level=1)
    document.add_paragraph(
        "报告不包含原始 PDF 或简历全文，常见电话、邮箱和详细地址会被脱敏。"
        "所有评分仅用于当前岗位之间的相对整理，不代表录取概率或 ATS 通过率。"
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_tailored_resume_docx(
    resume_profile: ResumeProfile,
    job_profile: JobProfile,
    accepted_suggestions: list[tuple[str, str]],
) -> bytes:
    """Create an editable resume content draft from parsed facts and accepted rewrites."""
    document = Document()
    _style_docx(document)
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    normal = document.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    document.styles["Title"].font.size = Pt(21)
    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)
    document.styles["Heading 1"].font.size = Pt(14)
    document.core_properties.title = "定制简历内容草稿"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_font(title.add_run("定制简历内容草稿"))
    target = document.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = target.add_run(f"目标岗位：{job_profile.company} · {job_profile.title}")
    run.bold = True
    _set_docx_font(run)

    if resume_profile.summary:
        document.add_heading("个人概要", level=1)
        paragraph = document.add_paragraph(redact_sensitive_info(resume_profile.summary))
        for run in paragraph.runs:
            _set_docx_font(run)

    if accepted_suggestions:
        document.add_heading("针对目标岗位的重点表述", level=1)
        for _, suggested_text in accepted_suggestions:
            _add_docx_bullet(document, suggested_text)

    if resume_profile.experience:
        document.add_heading("工作与实习经历", level=1)
        for item in resume_profile.experience:
            heading = document.add_paragraph()
            run = heading.add_run(
                " · ".join(filter(None, [item.organization, item.title]))
            )
            run.bold = True
            _set_docx_font(run)
            dates = " - ".join(filter(None, [item.start_date, item.end_date]))
            if dates or item.location:
                meta = document.add_paragraph(
                    " | ".join(filter(None, [dates, item.location]))
                )
                for meta_run in meta.runs:
                    _set_docx_font(meta_run)
            for bullet in item.bullets:
                _add_docx_bullet(document, bullet)

    if resume_profile.projects:
        document.add_heading("项目经历", level=1)
        for item in resume_profile.projects:
            heading = document.add_paragraph()
            run = heading.add_run(" · ".join(filter(None, [item.name, item.role])))
            run.bold = True
            _set_docx_font(run)
            for bullet in item.bullets:
                _add_docx_bullet(document, bullet)
            if item.technologies:
                _add_docx_bullet(document, "技术：" + "、".join(item.technologies))

    if resume_profile.education:
        document.add_heading("教育经历", level=1)
        for item in resume_profile.education:
            text = " · ".join(
                filter(None, [item.institution, item.degree, item.field_of_study])
            )
            _add_docx_bullet(document, text)
            for highlight in item.highlights:
                _add_docx_bullet(document, highlight)

    if resume_profile.skills:
        document.add_heading("技能", level=1)
        for group in resume_profile.skills:
            _add_docx_bullet(document, f"{group.category}：{'、'.join(group.skills)}")
    if resume_profile.languages:
        document.add_heading("语言", level=1)
        _add_docx_bullet(document, "、".join(resume_profile.languages))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_cover_letter_docx(
    job_profile: JobProfile,
    draft: CoverLetterDraft,
) -> bytes:
    """Create an editable cover letter from a validated evidence-bound draft."""
    document = Document()
    _style_docx(document)
    document.core_properties.title = "Cover Letter" if draft.language == "en" else "求职信"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""

    title_text = "Cover Letter" if draft.language == "en" else "求职信"
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_font(title.add_run(title_text))
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{job_profile.company} · {job_profile.title}")
    subtitle_run.bold = True
    _set_docx_font(subtitle_run)

    for text in [draft.salutation, *[item.text for item in draft.paragraphs], draft.closing]:
        paragraph = document.add_paragraph(redact_sensitive_info(text))
        for run in paragraph.runs:
            _set_docx_font(run)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_styles():
    if PDF_FONT_NAME not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, str(PDF_FONT_PATH)))
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ChineseTitle",
            parent=styles["Title"],
            fontName=PDF_FONT_NAME,
            fontSize=22,
            leading=30,
            alignment=TA_CENTER,
            textColor="#000000",
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "ChineseSubtitle",
            parent=styles["Normal"],
            fontName=PDF_FONT_NAME,
            fontSize=11,
            leading=17,
            alignment=TA_CENTER,
            textColor="#333333",
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "ChineseH1",
            parent=styles["Heading1"],
            fontName=PDF_FONT_NAME,
            fontSize=15,
            leading=21,
            textColor="#000000",
            spaceBefore=12,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "ChineseH2",
            parent=styles["Heading2"],
            fontName=PDF_FONT_NAME,
            fontSize=11.5,
            leading=17,
            textColor="#000000",
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "ChineseBody",
            parent=styles["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=10,
            leading=16,
            textColor="#111111",
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "ChineseBullet",
            parent=styles["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=9.5,
            leading=15,
            leftIndent=14,
            firstLineIndent=-8,
            textColor="#111111",
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "table": ParagraphStyle(
            "ChineseTable",
            parent=styles["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=9,
            leading=13,
            alignment=TA_CENTER,
            textColor="#111111",
            wordWrap="CJK",
        ),
        "table_left": ParagraphStyle(
            "ChineseTableLeft",
            parent=styles["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=9,
            leading=13,
            alignment=TA_LEFT,
            textColor="#111111",
            wordWrap="CJK",
        ),
        "table_header": ParagraphStyle(
            "ChineseTableHeader",
            parent=styles["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=9,
            leading=13,
            alignment=TA_CENTER,
            textColor="#FFFFFF",
            wordWrap="CJK",
        ),
    }


def _add_pdf_page_number(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont(PDF_FONT_NAME, 8)
    canvas.setFillColor("#555555")
    canvas.drawCentredString(letter[0] / 2, 0.45 * inch, f"第 {document.page} 页")
    canvas.restoreState()


def _add_landscape_pdf_page_number(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont(PDF_FONT_NAME, 8)
    canvas.setFillColor("#555555")
    canvas.drawCentredString(landscape(letter)[0] / 2, 0.38 * inch, f"第 {document.page} 页")
    canvas.restoreState()


def _pdf_paragraph(text: str, style) -> Paragraph:
    return Paragraph(escape(redact_sensitive_info(text)).replace("\n", "<br/>"), style)


def build_pdf_report(
    resume_profile: ResumeProfile,
    job_profile: JobProfile,
    analysis: MatchAnalysis,
    score: ScoreResult,
    interview_feedback: dict[str, dict] | None = None,
) -> bytes:
    """Create a paginated, in-memory PDF report with Chinese font support."""
    _ = resume_profile
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title="AI 求职分析报告",
        author="",
    )
    styles = _pdf_styles()
    story = [
        _pdf_paragraph("AI 求职分析报告", styles["title"]),
        _pdf_paragraph(f"{job_profile.company} · {job_profile.title}", styles["subtitle"]),
        _pdf_paragraph(
            f"工作地点：{job_profile.location or '未提供'}    岗位类型：{job_profile.job_type or '未提供'}",
            styles["subtitle"],
        ),
        _pdf_paragraph("核心结果", styles["h1"]),
        _pdf_paragraph(
            f"证据匹配度：{_score_text(score)}    信息完整度：{score.information_completeness:.1f}%    评分版本：{score.calculation_version}",
            styles["body"],
        ),
        _pdf_paragraph("岗位要求与证据", styles["h1"]),
    ]
    requirements = {item.id: item for item in job_profile.requirements}
    for item in analysis.matches:
        requirement = requirements[item.requirement_id]
        story.append(
            _pdf_paragraph(
                f"{requirement.normalized_name}  {STATUS_LABELS[item.status.value]}",
                styles["h2"],
            )
        )
        story.append(_pdf_paragraph(f"• JD：{requirement.original_text}", styles["bullet"]))
        story.append(_pdf_paragraph(f"• 判断：{item.explanation}", styles["bullet"]))
        evidence_text = "；".join(_evidence_lines(item)) or "无可验证证据"
        story.append(_pdf_paragraph(f"• 证据：{evidence_text}", styles["bullet"]))

    story.append(_pdf_paragraph("简历优化建议", styles["h1"]))
    if not analysis.resume_suggestions:
        story.append(_pdf_paragraph("本次没有通过证据校验的建议。", styles["body"]))
    for suggestion in analysis.resume_suggestions:
        story.append(
            _pdf_paragraph(
                f"• {_suggestion_source(analysis, suggestion.original_text)}：{suggestion.original_text}",
                styles["bullet"],
            )
        )
        story.append(_pdf_paragraph(f"• 建议：{suggestion.suggested_text}", styles["bullet"]))
        story.append(_pdf_paragraph(f"• 理由：{suggestion.reason}", styles["bullet"]))

    story.append(_pdf_paragraph("面试准备问题", styles["h1"]))
    if not analysis.interview_questions:
        story.append(_pdf_paragraph("本次没有生成可用的面试问题。", styles["body"]))
    for question in analysis.interview_questions:
        story.append(_pdf_paragraph(question.question, styles["h2"]))
        for point in question.answer_outline:
            story.append(_pdf_paragraph(f"• {point}", styles["bullet"]))

    feedback_items = interview_feedback or {}
    if feedback_items:
        story.append(_pdf_paragraph("模拟面试复盘", styles["h1"]))
        for record in feedback_items.values():
            feedback = record.get("feedback", {})
            story.append(_pdf_paragraph(record.get("question", "面试问题"), styles["h2"]))
            story.append(
                _pdf_paragraph(
                    "完整性 {}/5    STAR {}/5    岗位相关性 {}/5    表达清晰度 {}/5".format(
                        feedback.get("completeness_score", 0),
                        feedback.get("star_score", 0),
                        feedback.get("relevance_score", 0),
                        feedback.get("clarity_score", 0),
                    ),
                    styles["body"],
                )
            )
            for improvement in feedback.get("improvements", []):
                story.append(_pdf_paragraph(f"• 改进：{improvement}", styles["bullet"]))

    story.extend(
        [
            Spacer(1, 10),
            _pdf_paragraph("使用说明", styles["h1"]),
            _pdf_paragraph(
                "本报告基于用户提供的信息生成，不代表录取概率或 ATS 通过率。"
                "报告不包含原始简历全文，常见联系方式已脱敏。",
                styles["body"],
            ),
        ]
    )
    document.build(
        story,
        onFirstPage=_add_pdf_page_number,
        onLaterPages=_add_pdf_page_number,
    )
    return buffer.getvalue()


def build_job_comparison_pdf(items: list[JobComparisonItem]) -> bytes:
    """Create a landscape PDF comparison report with a readable ranking table."""
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(letter),
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.6 * inch,
        title="多岗位求职对比报告",
        author="",
    )
    styles = _pdf_styles()
    headers = ["序号", "公司与岗位", "分析与投递", "匹配度", "完整度", "ATS", "硬性风险", "必须缺口", "推荐值"]
    data = [[_pdf_paragraph(value, styles["table_header"]) for value in headers]]
    for rank, item in enumerate(items, start=1):
        data.append(
            [
                _pdf_paragraph(str(rank), styles["table"]),
                _pdf_paragraph(f"{item.company}\n{item.title}", styles["table_left"]),
                _pdf_paragraph(_comparison_stage_text(item), styles["table"]),
                _pdf_paragraph(_comparison_match_text(item), styles["table"]),
                _pdf_paragraph(f"{item.information_completeness:.1f}%", styles["table"]),
                _pdf_paragraph(str(item.ats_score), styles["table"]),
                _pdf_paragraph(str(item.hard_risks), styles["table"]),
                _pdf_paragraph(str(item.must_have_gaps), styles["table"]),
                _pdf_paragraph(f"{item.recommendation_score:.1f}", styles["table"]),
            ]
        )
    comparison_table = Table(
        data,
        colWidths=[0.45 * inch, 2.15 * inch, 0.9 * inch, 0.75 * inch, 0.8 * inch, 0.55 * inch, 0.75 * inch, 0.75 * inch, 0.7 * inch],
        repeatRows=1,
        hAlign="CENTER",
    )
    comparison_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#23415C")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, -1), PDF_FONT_NAME),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("ALIGN", (1, 1), (1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9D9D9")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F6FA")]),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story = [
        _pdf_paragraph("多岗位求职对比报告", styles["title"]),
        _pdf_paragraph(
            "汇总当前简历在各岗位下的本地计算结果。推荐值用于安排查看和投递顺序，"
            "不代表录取概率。",
            styles["subtitle"],
        ),
        _pdf_paragraph("岗位排序", styles["h1"]),
        comparison_table,
        Spacer(1, 12),
        _pdf_paragraph("隐私与使用说明", styles["h1"]),
        _pdf_paragraph(
            "报告不包含原始 PDF 或简历全文，常见电话、邮箱和详细地址会被脱敏。"
            "评分不代表任何招聘平台的 ATS 通过率。",
            styles["body"],
        ),
    ]
    document.build(
        story,
        onFirstPage=_add_landscape_pdf_page_number,
        onLaterPages=_add_landscape_pdf_page_number,
    )
    return buffer.getvalue()


def build_cover_letter_pdf(
    job_profile: JobProfile,
    draft: CoverLetterDraft,
) -> bytes:
    """Create a paginated PDF cover letter from a validated draft."""
    buffer = BytesIO()
    title_text = "Cover Letter" if draft.language == "en" else "求职信"
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.9 * inch,
        leftMargin=0.9 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.75 * inch,
        title=title_text,
        author="",
    )
    styles = _pdf_styles()
    story = [
        _pdf_paragraph(title_text, styles["title"]),
        _pdf_paragraph(f"{job_profile.company} · {job_profile.title}", styles["subtitle"]),
        Spacer(1, 12),
        _pdf_paragraph(draft.salutation, styles["body"]),
    ]
    for paragraph in draft.paragraphs:
        story.append(Spacer(1, 5))
        story.append(_pdf_paragraph(paragraph.text, styles["body"]))
    story.extend(
        [
            Spacer(1, 8),
            _pdf_paragraph(draft.closing, styles["body"]),
        ]
    )
    document.build(
        story,
        onFirstPage=_add_pdf_page_number,
        onLaterPages=_add_pdf_page_number,
    )
    return buffer.getvalue()


def build_markdown_report(
    resume_profile: ResumeProfile,
    job_profile: JobProfile,
    analysis: MatchAnalysis,
    score: ScoreResult,
    interview_feedback: dict[str, dict] | None = None,
) -> str:
    """Build an in-memory, redacted report without including the raw resume."""
    lines = [
        "# AI 求职分析报告",
        "",
        f"目标岗位：{job_profile.company} · {job_profile.title}",
        f"工作地点：{job_profile.location or '未提供'}",
        f"岗位类型：{job_profile.job_type or '未提供'}",
        "",
        "## 核心结果",
        "",
        f"- 证据匹配度：{'暂不可计算' if score.match_score is None else f'{score.match_score:.1f}%'}",
        f"- 信息完整度：{score.information_completeness:.1f}%",
        f"- 评分规则版本：{score.calculation_version}",
        "",
        "## 岗位要求与证据",
        "",
    ]
    requirements = {item.id: item for item in job_profile.requirements}
    for item in analysis.matches:
        requirement = requirements[item.requirement_id]
        lines.append(f"### {requirement.normalized_name}｜{STATUS_LABELS[item.status.value]}")
        lines.append("")
        lines.append(f"- JD：{requirement.original_text}")
        lines.append(f"- 判断：{item.explanation}")
        evidence_lines = _evidence_lines(item)
        if evidence_lines:
            lines.append("- 证据：" + "；".join(evidence_lines))
        else:
            lines.append("- 证据：无可验证证据")
        lines.append("")

    lines.extend(["## 简历优化建议", ""])
    if not analysis.resume_suggestions:
        lines.append("本次没有通过证据校验的建议。")
    for suggestion in analysis.resume_suggestions:
        lines.extend(
            [
                f"- {_suggestion_source(analysis, suggestion.original_text)}：{suggestion.original_text}",
                f"  建议：{suggestion.suggested_text}",
                f"  理由：{suggestion.reason}",
            ]
        )

    lines.extend(["", "## 面试准备问题", ""])
    for question in analysis.interview_questions:
        lines.append(f"- {question.question}")
        for point in question.answer_outline:
            lines.append(f"  - {point}")

    feedback_items = interview_feedback or {}
    if feedback_items:
        lines.extend(["", "## 模拟面试复盘", ""])
        for record in feedback_items.values():
            feedback = record.get("feedback", {})
            lines.append(f"### {record.get('question', '面试问题')}")
            lines.append("")
            lines.append(
                "- 四项评分：完整性 {}/5，STAR {}/5，岗位相关性 {}/5，表达清晰度 {}/5".format(
                    feedback.get("completeness_score", 0),
                    feedback.get("star_score", 0),
                    feedback.get("relevance_score", 0),
                    feedback.get("clarity_score", 0),
                )
            )
            for improvement in feedback.get("improvements", []):
                lines.append(f"- 改进：{improvement}")
            lines.append("")

    lines.extend(
        [
            "---",
            "本报告基于用户提供的信息生成，不代表录取概率或 ATS 通过率。",
            "报告不包含原始简历全文，常见联系方式已脱敏。",
        ]
    )
    return redact_sensitive_info("\n".join(lines)).strip() + "\n"
