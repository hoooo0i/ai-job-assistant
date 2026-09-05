from io import BytesIO

import fitz
from docx import Document

from src.reporting import (
    build_cover_letter_docx,
    build_cover_letter_pdf,
    build_docx_report,
    build_job_comparison_docx,
    build_job_comparison_pdf,
    build_markdown_report,
    build_pdf_report,
    build_tailored_resume_docx,
)
from src.schemas import (
    CoverLetterDraft,
    CoverLetterParagraph,
    JobProfile,
    JobComparisonItem,
    JobRequirement,
    MatchAnalysis,
    MatchEvidence,
    MatchStatus,
    RequirementCategory,
    RequirementImportance,
    RequirementMatch,
    ResumeProfile,
    ScoreResult,
)


def test_report_contains_results_but_redacts_contact_details() -> None:
    resume = ResumeProfile(
        summary="Synthetic candidate",
        education=[],
        experience=[],
        projects=[],
        skills=[],
        languages=[],
        evidence_chunks=[],
    )
    job = JobProfile(
        company="示例公司",
        title="数据岗位",
        location=None,
        job_type=None,
        responsibilities=[],
        requirements=[
            JobRequirement(
                id="req_001",
                original_text="需要 Python 经验",
                normalized_name="Python 经验",
                category=RequirementCategory.skill,
                importance=RequirementImportance.must_have,
                is_hard_condition=False,
            )
        ],
        domain_background=[],
    )
    analysis = MatchAnalysis(
        matches=[
            RequirementMatch(
                requirement_id="req_001",
                status=MatchStatus.matched,
                resume_evidence=["Python project, candidate@example.test, +61 412 345 678"],
                explanation="存在项目证据。",
                confidence=0.9,
            )
        ]
    )
    score = ScoreResult(
        match_score=100.0,
        information_completeness=100.0,
        known_weight=3,
        total_weight=3,
        calculation_version="v1.0",
    )

    report = build_markdown_report(resume, job, analysis, score)

    assert "示例公司 · 数据岗位" in report
    assert "证据匹配度：100.0%" in report
    assert "candidate@example.test" not in report
    assert "+61 412 345 678" not in report
    assert "[已隐藏邮箱]" in report

    docx_bytes = build_docx_report(resume, job, analysis, score)
    assert docx_bytes.startswith(b"PK")
    document = Document(BytesIO(docx_bytes))
    docx_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "AI 求职分析报告" in docx_text
    assert "candidate@example.test" not in docx_text
    assert "[已隐藏邮箱]" in docx_text

    pdf_bytes = build_pdf_report(resume, job, analysis, score)
    assert pdf_bytes.startswith(b"%PDF")
    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf:
        pdf_text = "\n".join(page.get_text() for page in pdf)
    assert "AI 求职分析报告" in pdf_text
    assert "candidate@example.test" not in pdf_text
    assert "[已隐藏邮箱]" in pdf_text


def test_tailored_resume_and_cover_letter_exports_are_readable() -> None:
    resume = ResumeProfile(
        summary="Data candidate",
        education=[],
        experience=[],
        projects=[],
        skills=[],
        languages=["English"],
        evidence_chunks=[],
    )
    job = JobProfile(
        company="示例公司",
        title="数据岗位",
        location=None,
        job_type=None,
        responsibilities=[],
        requirements=[],
        domain_background=[],
    )
    accepted = [("Original", "Tailored Python project statement.")]

    resume_bytes = build_tailored_resume_docx(resume, job, accepted)
    resume_document = Document(BytesIO(resume_bytes))
    resume_text = "\n".join(item.text for item in resume_document.paragraphs)
    assert "定制简历内容草稿" in resume_text
    assert "Tailored Python project statement." in resume_text

    draft = CoverLetterDraft(
        language="zh",
        salutation="尊敬的招聘团队：",
        paragraphs=[
            CoverLetterParagraph(
                text="我的项目经验与岗位要求相关。",
                evidence_ids=["ev_001"],
            )
        ],
        closing="感谢您的考虑。",
        caution_notes=[],
    )
    cover_docx = build_cover_letter_docx(job, draft)
    cover_document = Document(BytesIO(cover_docx))
    cover_text = "\n".join(item.text for item in cover_document.paragraphs)
    assert "求职信" in cover_text
    assert "项目经验" in cover_text

    cover_pdf = build_cover_letter_pdf(job, draft)
    with fitz.open(stream=cover_pdf, filetype="pdf") as pdf:
        pdf_text = "\n".join(page.get_text() for page in pdf)
    assert "求职信" in pdf_text
    assert "项目经验" in pdf_text


def test_report_labels_user_confirmed_evidence() -> None:
    resume = ResumeProfile(
        summary=None,
        education=[],
        experience=[],
        projects=[],
        skills=[],
        languages=[],
        evidence_chunks=[],
    )
    requirement = JobRequirement(
        id="req_001",
        original_text="需要 Python 经验",
        normalized_name="Python 经验",
        category=RequirementCategory.skill,
        importance=RequirementImportance.must_have,
        is_hard_condition=False,
    )
    job = JobProfile(
        company="示例公司",
        title="数据岗位",
        location=None,
        job_type=None,
        responsibilities=[],
        requirements=[requirement],
        domain_background=[],
    )
    analysis = MatchAnalysis(
        matches=[
            RequirementMatch(
                requirement_id="req_001",
                status=MatchStatus.partial,
                resume_evidence=[],
                evidence=[
                    MatchEvidence(
                        source="user_confirmed",
                        text="Built a Python workflow.",
                        fact_id="fact_001",
                    )
                ],
                explanation="用户补充了真实经历。",
                confidence=0.9,
            )
        ]
    )
    score = ScoreResult(
        match_score=50.0,
        information_completeness=100.0,
        known_weight=3,
        total_weight=3,
        calculation_version="v1.0",
    )

    report = build_markdown_report(resume, job, analysis, score)

    assert "[用户确认] Built a Python workflow." in report
    assert "[简历] Built a Python workflow." not in report


def test_multi_job_comparison_exports_are_readable() -> None:
    items = [
        JobComparisonItem(
            job_id="job_1",
            company="示例公司",
            title="数据分析师",
            stage="final",
            match_score=82.5,
            information_completeness=90.0,
            ats_score=88,
            hard_risks=0,
            must_have_gaps=1,
            recommendation_score=84.7,
        ),
        JobComparisonItem(
            job_id="job_2",
            company="第二家公司",
            title="AI 产品实习生",
            stage="clarification",
            match_score=None,
            information_completeness=40.0,
            ats_score=72,
            hard_risks=1,
            must_have_gaps=2,
            recommendation_score=5.4,
        ),
    ]

    docx_bytes = build_job_comparison_docx(items)
    document = Document(BytesIO(docx_bytes))
    docx_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "多岗位求职对比报告" in docx_text
    assert "示例公司" in docx_text
    assert "82.5%" in docx_text

    pdf_bytes = build_job_comparison_pdf(items)
    assert pdf_bytes.startswith(b"%PDF")
    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf:
        pdf_text = "\n".join(page.get_text() for page in pdf)
    assert "多岗位求职对比报告" in pdf_text
    assert "示例公司" in pdf_text
    assert "82.5%" in pdf_text
